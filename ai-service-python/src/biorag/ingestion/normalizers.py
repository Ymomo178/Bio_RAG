"""把常见文档格式转换为带图片资产引用的 Markdown。"""

import base64
import mimetypes
import re
from pathlib import Path
from urllib.parse import unquote, urlsplit

import pymupdf
from bs4 import BeautifulSoup
from docutils.core import publish_parts
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE
from markdownify import markdownify

from biorag.ingestion.models import ExtractedImage, NormalizedContent


VECTOR_DRAWING_THRESHOLD = 80
MINIMUM_PDF_IMAGE_SIDE = 50


def normalize_file(path: Path, fallback_title: str, asset_root: Path | None = None) -> NormalizedContent:
    """根据文件扩展名选择对应方法，并限制本地图片只从来源目录读取。"""
    path = path.resolve()
    asset_root = (asset_root or path.parent).resolve()
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _normalize_pdf(path, fallback_title)
    if suffix in {".html", ".htm"}:
        return _normalize_html(path, fallback_title, asset_root)
    if suffix == ".rst":
        return _normalize_rst(path, fallback_title, asset_root)
    if suffix in {".md", ".mdx"}:
        return _normalize_markdown(path, fallback_title, asset_root, is_mdx=suffix == ".mdx")
    if suffix == ".docx":
        return _normalize_docx(path, fallback_title)
    if suffix == ".txt":
        text = path.read_text(encoding="utf-8", errors="replace")
        return NormalizedContent(fallback_title, _clean_markdown(text))
    raise ValueError(f"暂不支持的文件格式：{path}")


def _normalize_pdf(path: Path, fallback_title: str) -> NormalizedContent:
    """使用 PyMuPDF 按页面坐标提取文本、表格、内嵌图片和重要矢量页面。"""
    document = pymupdf.open(path)
    metadata_title = document.metadata.get("title") if document.metadata else None
    title = _clean_inline_text(metadata_title) or fallback_title
    pages: list[str] = [f"# {title}"]
    images: list[ExtractedImage] = []

    try:
        for page_index, page in enumerate(document):
            page_number = page_index + 1
            page_blocks = _pdf_text_blocks(page)
            tables = _pdf_tables(page)
            page_items: list[tuple[float, float, str]] = []

            for block in page_blocks:
                bbox, text = block
                if text and not any(_overlap_ratio(bbox, table_bbox) >= 0.5 for table_bbox, _ in tables):
                    page_items.append((bbox[1], bbox[0], text))
            for table_bbox, table_markdown in tables:
                page_items.append((table_bbox[1], table_bbox[0], table_markdown))

            page_images = _extract_pdf_images(document, page, page_number, page_blocks, len(images))
            images.extend(page_images)
            for image in page_images:
                marker_y = float(image.section.split(":", maxsplit=1)[1]) if image.section and ":" in image.section else page.rect.height
                page_items.append((marker_y, 0.0, image.placeholder))

            if not page_images and len(page.get_drawings()) >= VECTOR_DRAWING_THRESHOLD:
                placeholder = _image_placeholder(len(images) + 1)
                pixmap = page.get_pixmap(matrix=pymupdf.Matrix(1.5, 1.5), alpha=False)
                vector_preview = ExtractedImage(
                    placeholder=placeholder,
                    filename=f"page-{page_number:04d}-vector-preview.png",
                    media_type="image/png",
                    content=pixmap.tobytes("png"),
                    caption=f"第 {page_number} 页页面预览（包含矢量图形）",
                    extraction_method="pdf-page-render-vector-fallback",
                    original_uri=None,
                    page_number=page_number,
                    section=f"第 {page_number} 页",
                    width=pixmap.width,
                    height=pixmap.height,
                )
                images.append(vector_preview)
                page_items.append((page.rect.height, 0.0, placeholder))

            ordered_content = "\n\n".join(item[2] for item in sorted(page_items))
            pages.append(f"## 第 {page_number} 页\n\n{ordered_content or '[本页没有可提取文本]'}")
    finally:
        document.close()
    return NormalizedContent(title, _clean_markdown("\n\n".join(pages)), tuple(images))


def _pdf_text_blocks(page: pymupdf.Page) -> list[tuple[tuple[float, float, float, float], str]]:
    """读取 PDF 文本块，并按照页面阅读顺序保留坐标。"""
    blocks: list[tuple[tuple[float, float, float, float], str]] = []
    for block in page.get_text("blocks", sort=True):
        if len(block) < 7 or block[6] != 0:
            continue
        text = _clean_markdown_fragment(block[4])
        blocks.append(((block[0], block[1], block[2], block[3]), text))
    return blocks


def _pdf_tables(page: pymupdf.Page) -> list[tuple[tuple[float, float, float, float], str]]:
    """检测页面表格并转换为 Markdown，检测失败时安全回退到普通文本。"""
    try:
        tables = page.find_tables().tables
    except (AttributeError, RuntimeError, ValueError):
        return []

    results: list[tuple[tuple[float, float, float, float], str]] = []
    for table in tables:
        rows = table.extract()
        markdown = _rows_to_markdown(rows)
        if markdown:
            results.append((tuple(table.bbox), markdown))
    return results


def _rows_to_markdown(rows: list[list[str | None]]) -> str:
    """把二维表格数据转换为 Markdown 表格。"""
    cleaned_rows = [
        [_clean_inline_text(cell).replace("|", "\\|") for cell in row]
        for row in rows
        if row
    ]
    if not cleaned_rows or not any(any(cell for cell in row) for row in cleaned_rows):
        return ""
    column_count = max(len(row) for row in cleaned_rows)
    normalized_rows = [row + [""] * (column_count - len(row)) for row in cleaned_rows]
    header = normalized_rows[0]
    body = normalized_rows[1:]
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join("---" for _ in range(column_count)) + " |",
    ]
    lines.extend("| " + " | ".join(row) + " |" for row in body)
    return "\n".join(lines)


def _extract_pdf_images(
    document: pymupdf.Document,
    page: pymupdf.Page,
    page_number: int,
    text_blocks: list[tuple[tuple[float, float, float, float], str]],
    existing_image_count: int,
) -> list[ExtractedImage]:
    """提取 PDF 页面中的原始位图，并记录所在页和附近图注。"""
    images: list[ExtractedImage] = []
    seen_occurrences: set[tuple[int, int, int, int, int]] = set()
    for image_info in page.get_images(full=True):
        xref = image_info[0]
        try:
            extracted = document.extract_image(xref)
            rectangles = page.get_image_rects(xref)
        except (RuntimeError, ValueError):
            continue
        if not extracted or not rectangles:
            continue
        for rectangle in rectangles:
            occurrence = (
                xref,
                round(rectangle.x0),
                round(rectangle.y0),
                round(rectangle.x1),
                round(rectangle.y1),
            )
            if occurrence in seen_occurrences:
                continue
            seen_occurrences.add(occurrence)
            if rectangle.width < MINIMUM_PDF_IMAGE_SIDE or rectangle.height < MINIMUM_PDF_IMAGE_SIDE:
                continue
            extension = extracted.get("ext", "png").lower()
            image_number = existing_image_count + len(images) + 1
            caption = _nearest_pdf_caption(rectangle, text_blocks) or f"第 {page_number} 页图片 {len(images) + 1}"
            images.append(
                ExtractedImage(
                    placeholder=_image_placeholder(image_number),
                    filename=f"page-{page_number:04d}-image-{len(images) + 1:03d}.{extension}",
                    media_type=_media_type_from_extension(extension),
                    content=extracted["image"],
                    caption=caption,
                    extraction_method="pdf-embedded-image",
                    original_uri=None,
                    page_number=page_number,
                    section=f"第 {page_number} 页:{rectangle.y0}",
                    width=extracted.get("width"),
                    height=extracted.get("height"),
                )
            )
    return images


def _nearest_pdf_caption(
    image_rectangle: pymupdf.Rect,
    text_blocks: list[tuple[tuple[float, float, float, float], str]],
) -> str | None:
    """从图片下方优先寻找距离最近的短文本作为图注。"""
    candidates: list[tuple[float, str]] = []
    for bbox, text in text_blocks:
        if not text or len(text) > 500:
            continue
        vertical_distance = bbox[1] - image_rectangle.y1
        horizontal_overlap = min(bbox[2], image_rectangle.x1) - max(bbox[0], image_rectangle.x0)
        if 0 <= vertical_distance <= 120 and horizontal_overlap > 0:
            candidates.append((vertical_distance, _clean_inline_text(text)))
    return min(candidates, default=(0, ""), key=lambda item: item[0])[1] or None


def _overlap_ratio(
    first: tuple[float, float, float, float],
    second: tuple[float, float, float, float],
) -> float:
    """计算第一个矩形被第二个矩形覆盖的面积比例。"""
    intersection_width = max(0.0, min(first[2], second[2]) - max(first[0], second[0]))
    intersection_height = max(0.0, min(first[3], second[3]) - max(first[1], second[1]))
    first_area = max(0.0, first[2] - first[0]) * max(0.0, first[3] - first[1])
    return (intersection_width * intersection_height / first_area) if first_area else 0.0


def _normalize_html(path: Path, fallback_title: str, asset_root: Path) -> NormalizedContent:
    """清除 HTML 页面框架，把正文和图片转换为 Markdown 与图片资产。"""
    html = path.read_text(encoding="utf-8", errors="replace")
    return _html_text_to_markdown(html, fallback_title, path.parent, asset_root)


def _html_text_to_markdown(
    html: str,
    fallback_title: str,
    base_directory: Path,
    asset_root: Path,
) -> NormalizedContent:
    """解析 HTML 字符串，保留正文结构、表格、图片图注和章节关系。"""
    soup = BeautifulSoup(html, "html.parser")
    root = (
        soup.select_one(".main-container")
        or soup.find("main")
        or soup.find("article")
        or soup.find("body")
        or soup
    )
    title_node = root.find("h1") if hasattr(root, "find") else None
    page_title = soup.find("title")
    title_text = title_node.get_text(" ", strip=True) if title_node else page_title.get_text(" ", strip=True) if page_title else ""
    title = _clean_inline_text(title_text) or fallback_title

    for selector in ("script", "style", "nav", "header", "footer", "form", "iframe", "noscript"):
        for node in root.select(selector):
            node.decompose()
    for selector in ("#TOC", ".tocify", ".toc-content", ".navbar"):
        for node in root.select(selector):
            node.decompose()

    images: list[ExtractedImage] = []
    for image in list(root.find_all("img")):
        source_uri = image.get("src") or ""
        alt_text = _clean_inline_text(image.get("alt") or image.get("title") or "")
        previous_heading = image.find_previous(["h1", "h2", "h3", "h4", "h5", "h6"])
        section = _clean_inline_text(previous_heading.get_text(" ", strip=True)) if previous_heading else None
        caption = _html_image_caption(image) or alt_text or (
            f"{section} 图" if section else f"图片 {len(images) + 1}"
        )
        asset = _asset_from_uri(
            source_uri=source_uri,
            caption=caption,
            section=section,
            image_number=len(images) + 1,
            base_directory=base_directory,
            asset_root=asset_root,
            extraction_prefix="html",
        )
        images.append(asset)
        image.replace_with(soup.new_string(asset.placeholder))

    converted = markdownify(str(root), heading_style="ATX", bullets="-")
    return NormalizedContent(title, _clean_markdown(converted), tuple(images))


def _html_image_caption(image: object) -> str:
    """从 figure 或相邻 caption 元素中读取 HTML 图片图注。"""
    figure = image.find_parent("figure")
    if figure:
        figure_caption = figure.find("figcaption")
        if figure_caption:
            return _clean_inline_text(figure_caption.get_text(" ", strip=True))
    parent = image.parent
    next_sibling = parent.find_next_sibling() if parent else None
    if next_sibling and "caption" in (next_sibling.get("class") or []):
        return _clean_inline_text(next_sibling.get_text(" ", strip=True))
    return ""


def _normalize_rst(path: Path, fallback_title: str, asset_root: Path) -> NormalizedContent:
    """使用 docutils 解析 RST，再复用 HTML 逻辑提取正文和图片。"""
    source = path.read_text(encoding="utf-8", errors="replace")
    source = _clean_sphinx_roles(source)
    parts = publish_parts(
        source=source,
        source_path=str(path),
        writer="html5",
        settings_overrides={
            "file_insertion_enabled": False,
            "raw_enabled": False,
            "report_level": 5,
            "halt_level": 6,
        },
    )
    return _html_text_to_markdown(parts["html_body"], fallback_title, path.parent, asset_root)


def _clean_sphinx_roles(source: str) -> str:
    """移除 docutils 无法解析的 Sphinx 角色，同时保留人类可读标签。"""
    def replace_role(match: re.Match[str]) -> str:
        """把 `显示文字 <目标>` 简化为显示文字。"""
        label = match.group(1)
        explicit_target = re.match(r"(.+?)\s*<[^>]+>\s*$", label, flags=re.DOTALL)
        return explicit_target.group(1).strip() if explicit_target else label.strip()

    return re.sub(r":[A-Za-z0-9_.-]+:`([^`]+)`", replace_role, source, flags=re.DOTALL)


def _normalize_markdown(
    path: Path,
    fallback_title: str,
    asset_root: Path,
    *,
    is_mdx: bool,
) -> NormalizedContent:
    """清理 Markdown/MDX，并把本地、内嵌和外部图片变成资产引用。"""
    text = path.read_text(encoding="utf-8", errors="replace")
    front_matter_title, text = _remove_front_matter(text)
    if is_mdx:
        text = re.sub(r"(?m)^\s*(?:import|export)\s+.*$", "", text)
        text = re.sub(r"</?[A-Z][A-Za-z0-9_.-]*\b[^>]*>", "", text)
    text = re.sub(r"<!--.*?-->", "", text, flags=re.DOTALL)

    images: list[ExtractedImage] = []
    current_section: str | None = None
    converted_lines: list[str] = []
    for line in text.splitlines():
        heading_match = re.match(r"^#{1,6}\s+(.+?)\s*$", line)
        if heading_match:
            current_section = _clean_inline_text(heading_match.group(1))

        def replace_markdown_image(match: re.Match[str]) -> str:
            """把一处 Markdown 图片语法替换为内部占位符。"""
            source_uri = match.group("uri").strip("<>")
            caption = _clean_inline_text(match.group("alt") or match.group("title") or "") or f"图片 {len(images) + 1}"
            asset = _asset_from_uri(
                source_uri=source_uri,
                caption=caption,
                section=current_section,
                image_number=len(images) + 1,
                base_directory=path.parent,
                asset_root=asset_root,
                extraction_prefix="markdown",
            )
            images.append(asset)
            return asset.placeholder

        line = re.sub(
            r"!\[(?P<alt>[^]]*)]\((?P<uri><[^>]+>|[^)\s]+)(?:\s+[\"'](?P<title>.*?)[\"'])?\)",
            replace_markdown_image,
            line,
        )

        def replace_html_image(match: re.Match[str]) -> str:
            """把 Markdown 中的 HTML img 标签替换为内部占位符。"""
            fragment = BeautifulSoup(match.group(0), "html.parser").find("img")
            source_uri = fragment.get("src") or ""
            caption = _clean_inline_text(fragment.get("alt") or fragment.get("title") or "") or f"图片 {len(images) + 1}"
            asset = _asset_from_uri(
                source_uri=source_uri,
                caption=caption,
                section=current_section,
                image_number=len(images) + 1,
                base_directory=path.parent,
                asset_root=asset_root,
                extraction_prefix="markdown-html",
            )
            images.append(asset)
            return asset.placeholder

        line = re.sub(r"<img\b[^>]*>", replace_html_image, line, flags=re.IGNORECASE)
        line = re.sub(r"<p\b[^>]*>\s*(BIORAGIMAGE\d{4}TOKEN)\s*</p>", r"\1", line, flags=re.IGNORECASE)
        converted_lines.append(line)

    text = "\n".join(converted_lines)
    title = front_matter_title or _first_markdown_heading(text) or fallback_title
    return NormalizedContent(title, _clean_markdown(text), tuple(images))


def _normalize_docx(path: Path, fallback_title: str) -> NormalizedContent:
    """读取 DOCX 段落、表格和媒体文件，并统一转换为 Markdown 与图片资产。"""
    document = Document(path)
    blocks: list[str] = []
    images: list[ExtractedImage] = []
    title = fallback_title
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = paragraph.style.name if paragraph.style else ""
        heading_match = re.match(r"Heading\s+(\d+)", style_name, flags=re.IGNORECASE)
        if heading_match:
            level = min(int(heading_match.group(1)), 6)
            blocks.append(f"{'#' * level} {text}")
            if level == 1:
                title = text
        else:
            blocks.append(text)
    for table in document.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        markdown_table = _rows_to_markdown(rows)
        if markdown_table:
            blocks.append(markdown_table)

    for relationship in document.part.rels.values():
        if relationship.reltype != RELATIONSHIP_TYPE.IMAGE:
            continue
        image_part = relationship.target_part
        image_number = len(images) + 1
        filename = Path(str(image_part.partname)).name
        images.append(
            ExtractedImage(
                placeholder=_image_placeholder(image_number),
                filename=filename,
                media_type=image_part.content_type,
                content=image_part.blob,
                caption=f"文档图片 {image_number}",
                extraction_method="docx-embedded-image",
                original_uri=f"word/media/{filename}",
                section="文档图片",
            )
        )
    if images:
        blocks.append("## 文档图片")
        blocks.extend(image.placeholder for image in images)
    return NormalizedContent(title, _clean_markdown("\n\n".join(blocks)), tuple(images))


def _asset_from_uri(
    source_uri: str,
    caption: str,
    section: str | None,
    image_number: int,
    base_directory: Path,
    asset_root: Path,
    extraction_prefix: str,
) -> ExtractedImage:
    """把 data URI、本地路径或外部 URL 转换为统一图片资产。"""
    placeholder = _image_placeholder(image_number)
    if source_uri.startswith("data:"):
        decoded = _decode_data_image(source_uri)
        if decoded:
            media_type, content, extension = decoded
            return ExtractedImage(
                placeholder=placeholder,
                filename=f"image-{image_number:04d}.{extension}",
                media_type=media_type,
                content=content,
                caption=caption,
                extraction_method=f"{extraction_prefix}-embedded-data-uri",
                original_uri="embedded:data-uri",
                section=section,
            )

    parsed = urlsplit(source_uri)
    if parsed.scheme in {"http", "https"}:
        filename = Path(unquote(parsed.path)).name or f"image-{image_number:04d}.bin"
        return ExtractedImage(
            placeholder=placeholder,
            filename=filename,
            media_type=mimetypes.guess_type(filename)[0] or "application/octet-stream",
            content=None,
            caption=caption,
            extraction_method=f"{extraction_prefix}-external-reference",
            original_uri=source_uri,
            section=section,
        )

    local_path = _resolve_local_asset(source_uri, base_directory, asset_root)
    if local_path:
        return ExtractedImage(
            placeholder=placeholder,
            filename=local_path.name,
            media_type=mimetypes.guess_type(local_path.name)[0] or "application/octet-stream",
            content=local_path.read_bytes(),
            caption=caption,
            extraction_method=f"{extraction_prefix}-local-file",
            original_uri=str(local_path),
            section=section,
        )

    filename = Path(unquote(parsed.path)).name or f"image-{image_number:04d}.bin"
    return ExtractedImage(
        placeholder=placeholder,
        filename=filename,
        media_type=mimetypes.guess_type(filename)[0] or "application/octet-stream",
        content=None,
        caption=caption,
        extraction_method=f"{extraction_prefix}-missing-reference",
        original_uri=source_uri or None,
        section=section,
    )


def _decode_data_image(source_uri: str) -> tuple[str, bytes, str] | None:
    """解码 Base64 图片 data URI，并返回媒体类型、内容和扩展名。"""
    match = re.match(r"^data:(image/[A-Za-z0-9.+-]+)(?:;[^,]*)?;base64,(.*)$", source_uri, flags=re.DOTALL)
    if not match:
        return None
    try:
        content = base64.b64decode(re.sub(r"\s+", "", match.group(2)), validate=False)
    except (ValueError, TypeError):
        return None
    media_type = match.group(1).lower()
    extension = {
        "image/jpeg": "jpg",
        "image/svg+xml": "svg",
        "image/tiff": "tiff",
    }.get(media_type, media_type.split("/", maxsplit=1)[1].replace("+xml", ""))
    return media_type, content, extension


def _resolve_local_asset(source_uri: str, base_directory: Path, asset_root: Path) -> Path | None:
    """在受限来源目录中解析本地图片，必要时按唯一文件名回退查找。"""
    path_text = unquote(urlsplit(source_uri).path)
    if not path_text:
        return None
    candidate = (base_directory / path_text).resolve()
    if _is_within(candidate, asset_root) and candidate.is_file():
        return candidate

    filename = Path(path_text).name
    if not filename:
        return None
    matches = [path for path in asset_root.rglob(filename) if path.is_file()]
    return matches[0].resolve() if len(matches) == 1 else None


def _is_within(path: Path, root: Path) -> bool:
    """判断解析后的文件是否仍位于允许的来源目录中。"""
    try:
        path.relative_to(root.resolve())
        return True
    except ValueError:
        return False


def _remove_front_matter(text: str) -> tuple[str | None, str]:
    """移除 Markdown 顶部 YAML 元数据，并尽量读取其中的标题。"""
    match = re.match(r"\A---\s*\n(.*?)\n---\s*\n?", text, flags=re.DOTALL)
    if not match:
        return None, text
    title_match = re.search(r"(?m)^title:\s*[\"']?(.*?)[\"']?\s*$", match.group(1))
    title = _clean_inline_text(title_match.group(1)) if title_match else None
    return title, text[match.end() :]


def _first_markdown_heading(text: str) -> str | None:
    """提取 Markdown 中第一个一级标题。"""
    match = re.search(r"(?m)^#\s+(.+?)\s*$", text)
    return _clean_inline_text(match.group(1)) if match else None


def _image_placeholder(image_number: int) -> str:
    """生成在流水线写出图片前使用的唯一占位符。"""
    return f"BIORAGIMAGE{image_number:04d}TOKEN"


def _media_type_from_extension(extension: str) -> str:
    """根据文件扩展名推断图片媒体类型。"""
    normalized = extension.lower().lstrip(".")
    if normalized == "jpg":
        return "image/jpeg"
    return mimetypes.guess_type(f"image.{normalized}")[0] or f"image/{normalized}"


def _clean_inline_text(text: str | None) -> str:
    """把标题等单行文本中的连续空白折叠为一个空格。"""
    return re.sub(r"\s+", " ", text or "").strip()


def _clean_markdown_fragment(text: str) -> str:
    """清理 PDF 文本块，同时保留块内必要换行。"""
    return "\n".join(line.rstrip() for line in text.replace("\r", "\n").splitlines()).strip()


def _clean_markdown(text: str) -> str:
    """统一换行、清除行尾空格，并确保图片编码不会混入正文。"""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"data:image/[^;\s]+;base64,[A-Za-z0-9+/=]+", "[内嵌图片已提取]", text)
    text = "\n".join(line.rstrip() for line in text.splitlines())
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip() + "\n"
